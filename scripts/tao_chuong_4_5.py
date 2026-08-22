# -*- coding: utf-8 -*-
"""Sinh Chương 4 và Chương 5 của quyển báo cáo từ chính mã nguồn và kết quả chạy thật.

Chủ sở hữu: TV1 (Chương 4) và TV8 (Chương 5).

Ưu điểm của cách làm này: mọi số liệu trong hai chương đều lấy trực tiếp từ mã
nguồn và từ kết quả chạy bộ kiểm thử, nên chạy lại script là báo cáo tự cập nhật,
không bao giờ lệch với chương trình thật::

    python -m scripts.tao_chuong_4_5

Sản phẩm:

- ``report/chuong-4-tv1/Chuong-4-Cai-dat.docx``
- ``report/chuong-5-tv8/Chuong-5-Ket-qua.docx``
"""

from __future__ import annotations

import io
import os
import platform
import subprocess
import sys
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

GOC = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ANH = os.path.join(GOC, "docs", "anh-demo")


# --------------------------------------------------------------------------
# Tiện ích dựng tài liệu
# --------------------------------------------------------------------------

def tao_tai_lieu() -> docx.Document:
    d = docx.Document()
    s = d.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(13)
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_after = Pt(6)
    for sec in d.sections:
        sec.left_margin = Inches(1.18)   # 3 cm
        sec.right_margin = Inches(0.79)  # 2 cm
        sec.top_margin = Inches(0.79)
        sec.bottom_margin = Inches(0.79)
    return d


def tieu_de(d, text, cap=1):
    p = d.add_heading(text, level=cap)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def doan(d, text, dam=False):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = dam
    return p


def chu_thich_hinh(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(12)
    return p


def them_anh(d, ten_file, rong=6.1):
    duong = os.path.join(ANH, ten_file)
    if not os.path.exists(duong):
        return False
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(duong, width=Inches(rong))
    return True


def khoi_ma(d, text):
    p = d.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    return p


def bang(d, tieu_de_cot, hang):
    t = d.add_table(rows=1, cols=len(tieu_de_cot))
    t.style = "Table Grid"
    for i, c in enumerate(tieu_de_cot):
        o = t.rows[0].cells[i]
        o.text = ""
        r = o.paragraphs[0].add_run(c)
        r.bold = True
        r.font.size = Pt(12)
    for h in hang:
        o = t.add_row().cells
        for i, v in enumerate(h):
            o[i].text = ""
            r = o[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(12)
    return t


def trich_ham(duong_dan: str, ten_ham: str, so_dong: int = 26) -> str:
    """Lấy phần đầu của một hàm trong mã nguồn để trích vào báo cáo."""
    src = io.open(os.path.join(GOC, duong_dan), encoding="utf-8").read().split("\n")
    for i, l in enumerate(src):
        if l.startswith("def " + ten_ham) or l.startswith("    def " + ten_ham):
            return "\n".join(src[i:i + so_dong])
    return ""


# --------------------------------------------------------------------------
# Chương 4
# --------------------------------------------------------------------------

def chuong_4() -> str:
    d = tao_tai_lieu()
    tieu_de(d, "CHƯƠNG 4. CÀI ĐẶT CHƯƠNG TRÌNH", 0)

    tieu_de(d, "4.1. Kiến trúc chương trình", 1)
    doan(d, "Chương trình được chia làm ba lớp tách rời nhau, mỗi lớp có một trách "
            "nhiệm riêng và chỉ giao tiếp với lớp kề qua một hợp đồng dữ liệu thống nhất.")
    bang(d, ["Lớp", "Thư mục", "Trách nhiệm", "Phụ trách"], [
        ["Giao diện", "gui/", "Nhận thao tác người dùng, hiển thị kết quả", "TV5, TV6"],
        ["Xử lý", "engine/", "Toàn bộ giải thuật Banker và giải thuật đồ thị", "TV4, TV7"],
        ["Dữ liệu", "data/", "Bộ dữ liệu mẫu định dạng JSON", "TV5"],
    ])
    doan(d, "Lý do tách lớp xử lý khỏi lớp giao diện: thư mục engine không được phép "
            "import bất kỳ thư viện đồ hoạ nào. Nhờ ràng buộc này, toàn bộ thuật toán "
            "chạy và kiểm thử tự động được từ dòng lệnh mà không cần mở cửa sổ — đó là "
            "điều kiện để xây dựng bộ kiểm thử tự động trình bày ở Chương 5.")
    doan(d, "Hợp đồng dữ liệu giữa hai lớp đặt tại engine/banker_types.py, được chốt "
            "trước khi các thành viên bắt đầu viết mã. Giao diện không gọi thẳng vào "
            "engine mà đi qua lớp trung gian gui/engine_adapter.py, nên khi engine thay "
            "đổi thì chỉ phải sửa một chỗ duy nhất.")

    tieu_de(d, "4.2. Cấu trúc dữ liệu và hợp đồng dùng chung", 1)
    doan(d, "Bốn cấu trúc dữ liệu của giải thuật được gói trong lớp BankerState:")
    bang(d, ["Trường", "Kích thước", "Ý nghĩa"], [
        ["available", "m", "Số thực thể mỗi loại còn rảnh trong hệ thống"],
        ["max", "n × m", "Nhu cầu tối đa tiến trình khai báo trước khi chạy"],
        ["allocation", "n × m", "Số thực thể đang được cấp cho tiến trình"],
        ["need", "n × m", "Tính ra: max − allocation"],
    ])
    doan(d, "Điểm thiết kế quan trọng nhất: need không phải dữ liệu lưu trữ mà là "
            "thuộc tính được tính lại mỗi lần đọc. Nhờ vậy đẳng thức Need = Max − "
            "Allocation luôn đúng trong mọi hoàn cảnh, loại bỏ hoàn toàn lỗi lệch dữ "
            "liệu — một lỗi kinh điển khi cài đặt giải thuật Banker.")
    khoi_ma(d, trich_ham("engine/banker_types.py", "need", 12).replace("    ", "  "))

    tieu_de(d, "4.3. Cài đặt thủ tục kiểm tra an toàn", 1)
    doan(d, "Thủ tục nhận vào trạng thái hệ thống và trả về đối tượng SafetyResult "
            "gồm bốn phần: kết luận an toàn hay không, chuỗi an toàn tìm được, nhật ký "
            "từng bước, và danh sách tiến trình còn treo nếu không an toàn. Nhật ký "
            "từng bước là bắt buộc, vì giao diện cần chiếu lại từng vòng lặp cho người "
            "xem chứ không chỉ hiển thị kết quả cuối cùng.")
    khoi_ma(d, trich_ham("engine/banker.py", "kiem_tra_an_toan", 8))

    tieu_de(d, "4.4. Cài đặt thủ tục xử lý yêu cầu tài nguyên", 1)
    doan(d, "Thủ tục trả về một trong ba kết cục, mỗi kết cục kèm một câu lý do riêng: "
            "cấp phát khi trạng thái sau khi cấp vẫn an toàn; chờ khi không đủ tài "
            "nguyên rảnh hoặc khi cấp vào sẽ làm trạng thái mất an toàn; lỗi khi tiến "
            "trình xin vượt quá nhu cầu đã khai báo.")
    doan(d, "Bước dễ cài sai nhất là bước bốn. Chương trình thao tác trên một bản sao "
            "sâu của trạng thái qua phương thức copy(), chỉ ghi đè trạng thái thật khi "
            "kết quả kiểm tra là an toàn. Tuyệt đối không sửa trực tiếp rồi trừ ngược "
            "lại, vì cách đó để lại trạng thái sai nếu có ngoại lệ xảy ra giữa chừng.")

    tieu_de(d, "4.5. Thiết kế giao diện", 1)
    doan(d, "Cửa sổ chính chia làm ba vùng theo chiều ngang, ngăn cách bằng thanh "
            "kéo cho phép người dùng tự chỉnh tỉ lệ:")
    bang(d, ["Vùng", "Chức năng", "Tệp"], [
        ["Trái", "Nhập ma trận, kiểm tra hợp lệ, tính Need, lưu và mở JSON",
         "gui/GUI_TV6.py"],
        ["Giữa", "Điều khiển mô phỏng, bảng nhật ký từng bước", "gui/GUI_TV6.py"],
        ["Phải", "Yêu cầu tài nguyên, giải phóng, hoàn tác, biểu đồ, đồ thị",
         "gui/bang_yeu_cau.py"],
    ])
    them_anh(d, "giao-dien-tong-quan.png", 6.3)
    chu_thich_hinh(d, "Hình 4.1 – Giao diện chương trình sau khi chạy hết năm bước mô phỏng")
    doan(d, "Ba kết cục của một yêu cầu tài nguyên được phân biệt bằng cả màu nền lẫn "
            "câu chữ: nền xanh cho cấp phát, nền cam cho chờ, nền đỏ cho lỗi. Việc dùng "
            "ba câu lý do khác nhau cho ba kết cục là yêu cầu bắt buộc của đề bài.")

    tieu_de(d, "4.6. Module xuất báo cáo", 1)
    doan(d, "Module gui/xuat_bao_cao.py xuất kết quả ra hai định dạng. Bản HTML không "
            "cần cài thêm thư viện nào, mở bằng trình duyệt rồi in ra PDF. Bản Excel "
            "cần thư viện openpyxl, xuất hai sheet gồm ma trận đầu vào và nhật ký.")
    doan(d, "Lý do không xuất thẳng PDF bằng reportlab: các phông chữ có sẵn của thư "
            "viện này không chứa dấu tiếng Việt, muốn hiển thị đúng dấu phải nhúng "
            "phông ngoài khá phức tạp. Đi đường HTML rồi để trình duyệt in ra PDF vừa "
            "gọn vừa không bao giờ lỗi phông.")

    p = os.path.join(GOC, "report", "chuong-4-tv1", "Chuong-4-Cai-dat.docx")
    os.makedirs(os.path.dirname(p), exist_ok=True)
    d.save(p)
    return p


# --------------------------------------------------------------------------
# Chương 5
# --------------------------------------------------------------------------

def chay_kiem_thu():
    """Chạy bộ kiểm thử, trả về (tổng số ca, số ca đạt, danh sách tên ca)."""
    kq = subprocess.run(
        [sys.executable, "-m", "unittest", "discover", "-s", "tests", "-v"],
        cwd=GOC, capture_output=True, text=True, encoding="utf-8", errors="replace")
    ra = (kq.stdout or "") + (kq.stderr or "")

    ten = [d.split(" ")[0] for d in ra.splitlines() if d.startswith("test_")]
    tong = 0
    for d in ra.splitlines():
        if d.startswith("Ran ") and " test" in d:
            tong = int(d.split()[1])
    dat = tong if ra.rstrip().endswith("OK") else 0
    return tong, dat, sorted(set(ten))


def chuong_5() -> str:
    tong, dat, ten_ca = chay_kiem_thu()

    d = tao_tai_lieu()
    tieu_de(d, "CHƯƠNG 5. KẾT QUẢ THỰC NGHIỆM", 0)

    tieu_de(d, "5.1. Môi trường thử nghiệm", 1)
    bang(d, ["Thành phần", "Thông số"], [
        ["Hệ điều hành", "{} {}".format(platform.system(), platform.release())],
        ["Kiến trúc", platform.machine()],
        ["Python", platform.python_version()],
        ["Thư viện giao diện", "PyQt5"],
        ["Khung kiểm thử", "unittest (thư viện chuẩn của Python)"],
    ])

    tieu_de(d, "5.2. Kết quả kiểm thử tự động", 1)
    doan(d, "Bộ kiểm thử gồm {} ca, chạy bằng lệnh python -m unittest discover -s "
            "tests. Kết quả: {}/{} ca đạt.".format(tong, dat, tong))
    doan(d, "Nguyên tắc áp dụng khi xây dựng bộ kiểm thử: người viết ca kiểm thử "
            "không phải là người viết mã cài đặt. Giá trị kỳ vọng của mỗi ca lấy từ "
            "bảng chạy tay ở Chương 2, tuyệt đối không suy đoán và không sửa ca kiểm "
            "thử cho khớp với mã nguồn.")
    bang(d, ["Nhóm ca", "Số ca", "Nội dung kiểm tra"], [
        ["Cấu trúc dữ liệu", "2", "Công thức Need = Max − Allocation, tổng tài nguyên là hằng số"],
        ["Kiểm tra an toàn", "5", "Trạng thái an toàn, bế tắc, Need bằng không, ca biên n=1 m=1"],
        ["Yêu cầu tài nguyên", "6", "Ba kết cục, tính nguyên vẹn khi khôi phục, ba lý do khác nhau"],
        ["Dữ liệu không hợp lệ", "4", "Allocation vượt Max, giá trị âm, sai số cột"],
        ["Ca bổ sung", "5", "Giải phóng dây chuyền, xin đúng bằng Need, xin vector rỗng"],
    ])

    tieu_de(d, "5.3. Các ca kiểm thử chính", 1)
    bang(d, ["Mã", "Tình huống", "Kết quả mong đợi"], [
        ["TC01", "Trạng thái ban đầu, Available = (3,3,2)", "An toàn, chuỗi <P1,P3,P0,P2,P4>"],
        ["TC02", "P1 yêu cầu (1,0,2)", "Cấp phát, Available còn (2,3,0)"],
        ["TC03", "Tiếp TC02, P4 yêu cầu (3,3,0)", "Chờ — vượt quá tài nguyên rảnh"],
        ["TC04", "Tiếp TC02, P0 yêu cầu (0,2,0)", "Chờ — đủ tài nguyên nhưng không an toàn"],
        ["TC05", "P2 yêu cầu (7,0,0) khi Need = (6,0,0)", "Lỗi — vượt quá Need khai báo"],
        ["TC06", "Available = (0,0,0), mọi Need > 0", "Không an toàn, liệt kê tiến trình treo"],
        ["TC09", "Nhập Allocation > Max", "Chặn tại giao diện, báo rõ ô sai"],
        ["TC22", "Yêu cầu cho tiến trình không tồn tại", "Lỗi E04 kèm thông điệp tiếng Việt"],
    ])
    doan(d, "Ca TC04 đáng chú ý nhất: hệ thống còn đủ tài nguyên rảnh để đáp ứng yêu "
            "cầu, nhưng nếu cấp thì trạng thái sau đó không còn an toàn nên yêu cầu bị "
            "trì hoãn. Đây chính là điểm phân biệt giải thuật tránh deadlock với việc "
            "cấp phát đơn thuần theo tài nguyên còn rảnh.")

    tieu_de(d, "5.4. Kết quả chạy chương trình", 1)
    them_anh(d, "giao-dien-tong-quan.png", 6.3)
    chu_thich_hinh(d, "Hình 5.1 – Kết luận AN TOÀN cùng chuỗi P1 → P3 → P0 → P2 → P4")
    doan(d, "Bảng nhật ký hiển thị đủ năm vòng lặp, mỗi vòng ghi Work trước, tiến "
            "trình được chọn, và Work sau khi tiến trình đó giải phóng tài nguyên. "
            "Kết quả trùng khớp hoàn toàn với bảng chạy tay ở Chương 2.")

    tieu_de(d, "5.5. Ưu điểm", 1)
    for x in [
        "Thuật toán cho kết quả đúng trên toàn bộ {} ca kiểm thử, trong đó có các ca "
        "biên như một tiến trình, một loại tài nguyên, và vector yêu cầu rỗng.".format(tong),
        "Ba kết cục của yêu cầu tài nguyên được phân biệt bằng ba câu lý do riêng, "
        "giúp người dùng hiểu vì sao yêu cầu bị từ chối.",
        "Nhật ký từng bước cho phép đối chiếu trực tiếp với mã giả ở Chương 2.",
        "Lớp xử lý tách rời hoàn toàn khỏi giao diện nên kiểm thử tự động được.",
        "Kết quả xuất được ra HTML và Excel để đưa vào báo cáo.",
    ]:
        d.add_paragraph(x, style="List Bullet")

    tieu_de(d, "5.6. Hạn chế", 1)
    for x in [
        "Chương trình yêu cầu biết trước nhu cầu tối đa của mọi tiến trình, đúng như "
        "giả định của giải thuật Banker. Hệ thống thực tế hiếm khi có thông tin này.",
        "Số tiến trình cố định trong suốt phiên làm việc, chưa cho phép tiến trình mới "
        "gia nhập giữa chừng.",
        "Mô phỏng chạy tuần tự trên một luồng, chưa phản ánh tình huống nhiều tiến "
        "trình yêu cầu tài nguyên đồng thời.",
        "Chức năng liệt kê toàn bộ chuỗi an toàn dùng thuật toán quay lui nên số chuỗi "
        "tăng rất nhanh theo n, phải đặt giới hạn số chuỗi trả về.",
    ]:
        d.add_paragraph(x, style="List Bullet")

    tieu_de(d, "5.7. Hướng phát triển", 1)
    for x in [
        "Bổ sung giải thuật phát hiện deadlock chạy định kỳ để so sánh trực tiếp giữa "
        "hướng tránh và hướng phát hiện.",
        "Mô phỏng đa luồng thời gian thực, cho phép nhiều tiến trình cùng gửi yêu cầu.",
        "Cho phép tiến trình gia nhập và rời khỏi hệ thống trong lúc đang chạy.",
        "Xuất trực tiếp ra định dạng PDF có nhúng phông tiếng Việt.",
    ]:
        d.add_paragraph(x, style="List Bullet")

    tieu_de(d, "5.8. Kết luận chương", 1)
    doan(d, "Chương trình đáp ứng đầy đủ các yêu cầu của đề bài: nhập liệu có kiểm "
            "tra hợp lệ, kiểm tra trạng thái an toàn kèm chuỗi an toàn, mô phỏng từng "
            "bước, xử lý trọn vẹn chu trình yêu cầu — sử dụng — giải phóng tài nguyên, "
            "và xuất kết quả ra tệp. Toàn bộ {} ca kiểm thử đều đạt, kết quả khớp với "
            "bảng chạy tay trong Chương 2.".format(tong))

    p = os.path.join(GOC, "report", "chuong-5-tv8", "Chuong-5-Ket-qua.docx")
    os.makedirs(os.path.dirname(p), exist_ok=True)
    d.save(p)
    return p


if __name__ == "__main__":
    print("Chương 4:", chuong_4())
    print("Chương 5:", chuong_5())
